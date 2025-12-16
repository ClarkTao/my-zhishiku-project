"""
=== Python代码文件: tender_engine.py ===
"""
import os
import re
import json
import traceback
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser, StrOutputParser
from bs4 import BeautifulSoup  # ✅ [新增] 用于解析 HTML 表格


# ==============================================================================
# 模块 1: 智能文档解析器 (DocParser - 增强版) (无修改)
# ==============================================================================
class DocParser:
    @staticmethod
    def parse_docx_structure(file_path: str) -> Dict[str, str]:
        """
        解析 Word 文档，返回 {章节标题: 章节内容} 的字典。
        支持复杂的正则匹配，确保不错过任何章节。
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = Document(file_path)
        content_map = {}
        current_header = "前言/未分类内容"
        current_text = []

        # 匹配：第x章、1.1.1、一、(1) 等
        header_pattern = re.compile(
            r'^\s*('
            r'第[零一二三四五六七八九十百]+[章节]|'  # 第一章
            r'[零一二三四五六七八九十百]+[、\.]|'  # 一、
            r'\d+(\.\d+)*|'  # 1.1.1
            r'\(\d+\)|'  # (1)
            r'（\d+）'  # （1）
            r')\s*.*'
        )

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 判断是否为标题：样式为 Heading 开头，或者符合正则且长度适中
            is_style_heading = para.style.name.startswith('Heading')
            is_regex_heading = header_pattern.match(text) and len(text) < 50

            if is_style_heading or is_regex_heading:
                # 保存上一章节的内容
                if current_text:
                    content_map[current_header] = "\n".join(current_text)

                # 开启新章节
                current_header = text
                current_text = []
            else:
                current_text.append(text)

        # 保存最后一个章节
        if current_text:
            content_map[current_header] = "\n".join(current_text)

        return content_map


# ==============================================================================
# 模块 2: 核心写作引擎 (TenderWriterEngine - 全功能版)
# ==============================================================================
class TenderWriterEngine:
    def __init__(self, api_key: str,
                 writer_model: str = "deepseek-chat",  # 写作：用聪明的大模型
                 auditor_model: str = "deepseek-chat"):  # 审核：未来可换成更快的 deepseek-lite 或其他小模型
        # 1. 写作模型
        self.llm = ChatOpenAI(
            model_name=writer_model,
            openai_api_key=api_key,
            openai_api_base="https://api.deepseek.com",
            temperature=0.7,  # 写作要更有有创意
            max_tokens=4000,
            request_timeout=600,  # ✅ [新增] 设置超时为 600秒 (10分钟)，防止长章节生成中断
            max_retries=3  # ✅ [新增] 失败自动重试 3 次
        )
        # 2. 审核员模型
        self.auditor_llm = ChatOpenAI(
            model_name=auditor_model,
            openai_api_key=api_key,
            openai_api_base="https://api.deepseek.com",
            temperature=0.1,  # 审核要严谨，不要发散
            max_tokens=4000,
            request_timeout=300,  # ✅ [新增] 同样设置超时
            max_retries=3
        )
        self.ref_content_map = {}
        self.ref_toc = []

    def load_reference(self, file_path: str):
        """加载参考标书"""
        print(f"📖 [系统] 正在深度解析参考文档: {file_path}...")
        try:
            self.ref_content_map = DocParser.parse_docx_structure(file_path)
            self.ref_toc = list(self.ref_content_map.keys())
            print(f"✅ [系统] 参考文档加载完毕，提取到 {len(self.ref_toc)} 个章节。")
        except Exception as e:
            print(f"❌ [错误] 解析文档失败: {e}")
            traceback.print_exc()
            raise e

    def analyze_style(self) -> str:
        """提取文风 DNA"""
        print("🧬 [AI] 正在提取参考文档的文风 DNA...")
        if not self.ref_content_map:
            return "专业、严谨、符合工程标书规范"

        sample_text = ""
        # 取前3个非空章节作为样本
        count = 0
        for k, v in self.ref_content_map.items():
            if len(v) > 50:
                sample_text += v + "\n"
                count += 1
            if count >= 3: break

        sample_text = sample_text[:2000]

        prompt = ChatPromptTemplate.from_template("""
        你是一位资深的文案风格分析师。请分析以下工程标书片段的写作风格：

        【片段】
        {text}

        【任务】
        请提取该文档的“文风 DNA”，简要总结以下三点（100字以内）：
        1. 语气基调（如：极度自信、客观中立、侧重技术细节）。
        2. 句式特征（如：多用排比短句、多用“必须/确保”等强硬词汇）。
        3. 核心术语习惯。

        返回一段简短的 Prompt 指令，例如：“请保持客观严谨的语气，多使用无主句，强调‘安全第一’...”
        """)

        try:
            chain = prompt | self.llm | StrOutputParser()
            style_dna = chain.invoke({"text": sample_text})
            print(f"🧬 [文风 DNA] {style_dna}")
            return style_dna
        except Exception as e:
            print(f"⚠️ 文风分析失败: {e}")
            return "保持专业工程标书风格，语言严谨规范。"

    # --- Step 1: 智能目录生成 ---
    def generate_outline(self, target_project_info: str) -> List[str]:
        prompt = ChatPromptTemplate.from_template("""
        你是一位拥有20年经验的标书主笔。

        【参考标书目录】
        {ref_toc}

        【新项目关键信息】
        {project_info}

        【任务】
        请为新项目设计一份标准的标书目录结构（List）。
        要求：
        1. 逻辑复用：尽量沿用参考标书的成熟框架。
        2. 针对性调整：如果新项目是“隧道”，参考标书是“桥梁”，请修改相关技术章节的标题。
        3. 格式统一：保持“1. xxx”, “1.1 xxx”的层级格式。
        4. 纯净输出：只返回一个 JSON 字符串列表，不要包含 Markdown 标记。
        """)

        chain = prompt | self.llm | JsonOutputParser()
        print("🤔 [AI] 正在构思新标书大纲...")
        return chain.invoke({
            "ref_toc": "\n".join(self.ref_toc[:50]),
            "project_info": target_project_info
        })

    # --- Step 2: 目录智能映射 (分批处理) ---
    def map_toc_relationships(self, new_toc: List[str]) -> Dict[str, str]:
        print("🔗 [系统] 正在以分批模式建立新旧章节映射关系...")
        final_mapping = {}
        batch_size = 15
        num_batches = (len(new_toc) + batch_size - 1) // batch_size
        ref_toc_subset_str = "\n".join(self.ref_toc[:40])

        for i in range(num_batches):
            print(f"  -> 正在处理批次 {i + 1}/{num_batches}...")
            batch_new_toc = new_toc[i * batch_size: (i + 1) * batch_size]
            batch_new_toc_str = "\n".join(batch_new_toc)

            prompt_text = f"""
               你是一个精准的目录匹配助手。
               【参考目录 (Source)】
               {ref_toc_subset_str}
               【待匹配目录 (Target - Current Batch)】
               {batch_new_toc_str}
               【任务】
               请为 "待匹配目录" 中的每一项，从 "参考目录" 中找出一个语义最相似的标题。
               【输出格式】
               返回一个 JSON 对象，Key 是 "待匹配目录" 的标题，Value 是 "参考目录" 中最相似的标题。
               如果找不到任何相似的，Value 设为 null。
               只返回纯净的 JSON 对象，不要包含其他任何解释或 Markdown 标记。
               """
            try:
                # 临时关闭流式以确保 JSON 完整
                original_streaming = self.llm.streaming
                self.llm.streaming = False
                response = self.llm.invoke(prompt_text)
                self.llm.streaming = original_streaming

                # 清洗 Markdown
                content = response.content.replace("```json", "").replace("```", "").strip()
                batch_mapping = json.loads(content)
                final_mapping.update(batch_mapping)
            except Exception as e:
                print(f"⚠️ [警告] 批次 {i + 1} 目录映射失败: {e}")
                for title in batch_new_toc:
                    final_mapping[title] = None

        print("✅ [系统] 所有批次的目录映射完成。")
        return final_mapping

    # ✅ [新增]：私有方法 - "审核员 Agent"
    def _review_and_fix(self, draft_content: str, project_info: str, style_guide: str) -> str:
        """
        审核员逻辑：检查草稿是否符合项目要求，是否存在机器幻觉，并修正。
        """
        print("🧐 [审核员] 正在校验内容...")
        audit_prompt = ChatPromptTemplate.from_template("""
        你是一位极其严格的标书审核专家（Auditor）。你的任务是审查并修正下属提交的草稿。

        【新项目真实背景】
        {project_info}

        【文风要求 (Style DNA)】
        {style_guide}

        【待审核的草稿内容】
        {draft}

        【审核清单】
        1. **内容准确性**：是否出现了参考范文中的旧地名、旧数据，而未替换为新项目的？
        2. **逻辑一致性**：内容是否与新项目背景完全相符？
        3. **风格一致性**：内容是否严格遵循了指定的文风要求？
        4. **机器幻觉**：是否存在明显的、无根据的胡编乱造？

        【输出指令】
        - 如果草稿质量非常高（95分以上），可以直接原样输出，不要添加任何额外文字。
        - 如果存在任何问题，请**直接输出修正后的、完美的最终版本**。不要解释你修改了哪里，也不要说“已修正”等任何废话。
        """)

        chain = audit_prompt | self.auditor_llm | StrOutputParser()
        return chain.invoke({
            "project_info": project_info,
            "style_guide": style_guide,
            "draft": draft_content
        })

    # --- Step 3: 深度仿写 ---
    # ✅ [修改]：修改了 Prompt 以强制输出 HTML 表格
    def write_chapter(self, chapter_title: str, ref_chapter_title: Optional[str], project_info: str,
                      style_guide: str = "") -> str:
        """
        生成单章内容，并引入“生成-校验-修正”闭环。
        """
        ref_content = ""
        if ref_chapter_title and ref_chapter_title in self.ref_content_map:
            ref_content = self.ref_content_map[ref_chapter_title][:3000]

        prompt_template = """
        你现在的身份是标书撰写专家。请编写章节：【{title}】。

        【新项目背景】
        {project_info}
        """

        # 注入文风要求
        if style_guide:
            prompt_template += f"""
        【文风要求 (Style DNA)】
        {style_guide}
            """

        if ref_content:
            prompt_template += """
        【参考范文 (严格模仿其格式、语气和术语)】
        {ref_content}

        【写作指令】
        1. **深度模仿**：结构、语气和分点方式必须与范文一致。
        2. **数据置换**：严禁保留范文中的旧地名、旧参数，必须全部替换为新项目数据。
        3. **逻辑修正**：确保技术描述符合新项目实际情况。
            """
        else:
            prompt_template += """
        【写作指令】
        未找到参考范文，请凭借专业知识，根据新项目背景和上述文风要求直接撰写。
            """

        # ✅ [新增] 强制 HTML 表格指令
        prompt_template += """
        ⚠️【重要排版指令 - 必须遵守】⚠️
        1. 正文内容请使用标准的 Markdown 格式（如 # 标题，- 列表）。
        2. 如果内容中包含表格（例如：工程量清单、设备表、人员表），**必须且只能**使用 HTML `<table>` 代码格式输出！
           - 严禁使用 Markdown 表格（|---|），因为它无法处理合并单元格。
           - 必须正确使用 `<thead>`, `<tbody>`, `<tr>`, `<td>` 标签。
           - 如需合并单元格，请务必使用 `rowspan` 和 `colspan` 属性。
           - 表格代码不需要包裹在 ```html ``` 代码块中，直接输出即可。
        """

        prompt = ChatPromptTemplate.from_template(prompt_template)
        chain = prompt | self.llm | StrOutputParser()

        # 1. 生成初稿
        print(f"✍️ [写手] 正在撰写章节: {chapter_title}...")
        draft_content = chain.invoke({
            "title": chapter_title,
            "project_info": project_info,
            "ref_content": ref_content
        })

        # 2. 审核与修正
        final_content = self._review_and_fix(draft_content, project_info, style_guide)

        return final_content

    # ✅ [新增] 核心算法：HTML 表格渲染器
    def _process_html_table(self, doc, html_content):
        """
        将 HTML 表格 (含 rowspan/colspan) 完美还原到 Word 文档中
        """
        soup = BeautifulSoup(html_content, 'html.parser')
        rows = soup.find_all('tr')
        if not rows: return

        # 1. 预计算表格维度
        n_rows = len(rows)
        n_cols = 0
        for tr in rows:
            current_row_cols = 0
            for cell in tr.find_all(['td', 'th']):
                current_row_cols += int(cell.get('colspan', 1))
            n_cols = max(n_cols, current_row_cols)

        if n_cols == 0: return

        # 2. 创建 Word 表格
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.style = 'Table Grid'

        # 3. 建立网格占用图
        grid_map = [[False for _ in range(n_cols)] for _ in range(n_rows)]

        # 4. 填充数据并执行合并
        for r_idx, tr in enumerate(rows):
            c_idx = 0
            cells = tr.find_all(['td', 'th'])

            for cell in cells:
                # 跳过已被占用的格子
                while c_idx < n_cols and grid_map[r_idx][c_idx]:
                    c_idx += 1

                if c_idx >= n_cols: break

                # 获取属性
                rowspan = int(cell.get('rowspan', 1))
                colspan = int(cell.get('colspan', 1))
                text = cell.get_text(strip=True)

                # 填入内容
                try:
                    word_cell = table.cell(r_idx, c_idx)
                    word_cell.text = text
                except IndexError:
                    pass

                # 执行合并
                if rowspan > 1 or colspan > 1:
                    end_r = r_idx + rowspan - 1
                    end_c = c_idx + colspan - 1
                    if end_r < n_rows and end_c < n_cols:
                        try:
                            word_cell.merge(table.cell(end_r, end_c))
                        except:
                            pass

                # 标记占用
                for i in range(r_idx, r_idx + rowspan):
                    for j in range(c_idx, c_idx + colspan):
                        if i < n_rows and j < n_cols:
                            grid_map[i][j] = True

    # --- Step 4: 组装导出 ---
    # ✅ [重写] 改为混合解析模式
    def compile_to_word(self, toc_content_map: Dict[str, str], output_path: str):
        doc = Document()
        # 设置中文字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.add_heading('投标文件', 0)

        # 预编译正则，用于分离 HTML 表格
        # 匹配 <table>...</table>，包含换行，忽略大小写
        table_pattern = re.compile(r'(<table>.*?</table>)', re.DOTALL | re.IGNORECASE)

        for title, content in toc_content_map.items():
            # 添加章节标题
            doc.add_heading(title, level=1)

            if not content: continue

            # 1. 使用正则切分：分离 HTML 表格块与普通文本
            parts = table_pattern.split(content)

            for part in parts:
                if not part.strip():
                    continue

                # A. 如果是 HTML 表格块
                if part.strip().lower().startswith('<table>'):
                    try:
                        self._process_html_table(doc, part)
                        doc.add_paragraph("")  # 表格后加空行
                    except Exception as e:
                        print(f"⚠️ 表格渲染失败: {e}")
                        # 降级处理：直接作为文本写入，避免丢失信息
                        doc.add_paragraph(part)

                # B. 如果是普通 Markdown 文本 (复用原有的 Markdown 解析逻辑)
                else:
                    lines = part.strip().split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue

                        # 识别 Markdown 标题 (### )
                        if line.startswith('### '):
                            doc.add_heading(line.replace('### ', ''), level=3)
                        elif line.startswith('## '):
                            doc.add_heading(line.replace('## ', ''), level=2)
                        # 识别列表 (- )
                        elif line.startswith('- ') or line.startswith('* '):
                            p = doc.add_paragraph(line[2:], style='List Bullet')
                            p.paragraph_format.line_spacing = 1.5
                        # 识别数字列表 (1. )
                        elif re.match(r'^\d+\.\s', line):
                            p = doc.add_paragraph(line, style='List Number')
                            p.paragraph_format.line_spacing = 1.5
                        # 普通段落，处理粗体 (**text**)
                        else:
                            p = doc.add_paragraph()
                            # 简单的粗体解析逻辑
                            parts = re.split(r'(\*\*.*?\*\*)', line)
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    run = p.add_run(part[2:-2])
                                    run.bold = True
                                else:
                                    p.add_run(part)
                            p.paragraph_format.line_spacing = 1.5

        doc.save(output_path)
        print(f"💾 文件已保存: {output_path}")
