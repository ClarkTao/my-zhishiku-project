"""
ingestion/tender_parser.py
功能：集成 语义切分 + 表格语义化 + 深度XML扫描 + OCR图片识别 (RapidOCR)
"""
import os
import uuid
import re
import pdfplumber
import numpy as np
from typing import List, Dict
from dataclasses import dataclass
from docx import Document
from PIL import Image

# --- 1. OCR 模块引入 (新增) ---
try:
    from rapidocr_onnxruntime import RapidOCR
    # 初始化 OCR 引擎 (只初始化一次，自动下载模型)
    # det: 文本检测, rec: 文本识别
    ocr_engine = RapidOCR()
    HAS_OCR = True
    print("👀 [Parser] RapidOCR 引擎加载成功")
except ImportError:
    HAS_OCR = False
    print("⚠️ 未安装 rapidocr_onnxruntime，无法识别图片内容")

# --- 2. LangChain 组件导入 ---
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_experimental.text_splitter import SemanticChunker

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except ImportError:
        print("❌ 严重错误: 缺少 langchain-text-splitters 库")
        class RecursiveCharacterTextSplitter:
            def __init__(self, **kwargs): pass
            def split_text(self, text): return [text]

# --- 3. 自定义处理器导入 ---
try:
    from ingestion.processors import TableSummarizer, TableProcessor
except ImportError:
    TableSummarizer = None
    class TableProcessor:
        @staticmethod
        def table_to_markdown(data): return str(data)

@dataclass
class IndexableChunk:
    content: str
    metadata: Dict
    chunk_id: str
    parent_id: str = None
    is_parent: bool = False

class TenderDocParser:
    def __init__(self, project_info: Dict[str, str], use_advanced_mode: bool = True):
        self.project_info = project_info or {}
        self.use_advanced_mode = use_advanced_mode
        self.table_summarizer = TableSummarizer() if (TableSummarizer and use_advanced_mode) else None

        if self.use_advanced_mode:
            print("⏳ [Parser] 初始化语义切分模型 (BGE)...")
            self.embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-small-zh-v1.5")
            self.splitter = SemanticChunker(self.embeddings, breakpoint_threshold_type="percentile")
        else:
            print("🚀 [Parser] 使用快速递归切分模式...")
            self.splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)

    def parse_file(self, file_path: str) -> List[IndexableChunk]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self._parse_docx(file_path)
        else:
            print(f"⚠️ 跳过不支持的文件类型: {ext}")
            return []

    # --- OCR 辅助函数 (新增) ---
    def _run_ocr_on_image(self, img_obj) -> str:
        """输入 PIL Image，返回识别文字"""
        if not HAS_OCR: return ""
        try:
            # RapidOCR 需要 numpy array 格式
            img_np = np.array(img_obj)
            # result 结构: List[ [box_coord, text, confidence] ]
            result, _ = ocr_engine(img_np)
            if result:
                # 提取识别到的文本，过滤低置信度 (<0.6)
                txts = [line[1] for line in result if float(line[2]) > 0.6]
                return "\n".join(txts)
        except Exception as e:
            print(f"⚠️ OCR 识别出错: {e}")
        return ""

    def _deep_scan_docx(self, doc) -> str:
        """
        Word 深度 XML 扫描 (针对文本框)
        """
        try:
            xml = doc._element.xml
            text_nodes = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
            return "\n".join(text_nodes)
        except Exception as e:
            print(f"❌ 深度扫描失败: {e}")
            return ""

    def _parse_docx(self, file_path: str) -> List[IndexableChunk]:
        print(f"📝 [Parser] 处理 Word: {os.path.basename(file_path)}")
        all_chunks = []
        full_text_buffer = ""

        try:
            doc = Document(file_path)

            # 1. 尝试提取表格
            for table in doc.tables:
                data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                content = ""
                if self.use_advanced_mode and self.table_summarizer:
                    content = self.table_summarizer.summarize_table(data)
                else:
                    content = TableProcessor.table_to_markdown(data)

                if content and len(content) > 10:
                    all_chunks.append(IndexableChunk(
                        content=content,
                        metadata={**self.project_info, "source_file": os.path.basename(file_path), "page": 1, "type": "table"},
                        chunk_id=str(uuid.uuid4())
                    ))

            # 2. 尝试标准提取文本 (段落)
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text_buffer += para.text + "\n"

            # 3. 深度 XML 扫描 (文本框)
            if not full_text_buffer.strip():
                print("⚠️ [Parser] 标准解析未发现文本，启动深度 XML 扫描(文本框)...")
                full_text_buffer = self._deep_scan_docx(doc)
                if full_text_buffer:
                    print(f"✅ [Parser] 深度扫描成功提取 {len(full_text_buffer)} 字符")

            # 4. 纯图片 Word 提示
            if not full_text_buffer.strip() and not all_chunks:
                print("⚠️ [Parser] Word 文件似乎是纯图片，Word OCR 提取极其不稳定，建议用户转 PDF 上传。")
                return []

            # 5. 切分文本
            if full_text_buffer:
                text_chunks = self._split_text(full_text_buffer)
                for txt in text_chunks:
                    if len(txt.strip()) > 5:
                        all_chunks.append(IndexableChunk(
                            content=txt,
                            metadata={**self.project_info, "source_file": os.path.basename(file_path), "page": 1, "type": "text"},
                            chunk_id=str(uuid.uuid4())
                        ))

        except Exception as e:
            print(f"❌ Word 解析失败: {e}")

        return all_chunks

    def _parse_pdf(self, file_path: str) -> List[IndexableChunk]:
        print(f"📄 [Parser] 处理 PDF: {os.path.basename(file_path)}")
        all_chunks = []
        full_text_buffer = ""

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text_buffer = ""

                # A. 提取表格
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        content = ""
                        if self.use_advanced_mode and self.table_summarizer:
                            content = self.table_summarizer.summarize_table(table)
                        else:
                            content = TableProcessor.table_to_markdown(table)
                        if content:
                            all_chunks.append(IndexableChunk(
                                content=content,
                                metadata={**self.project_info, "source_file": os.path.basename(file_path), "page": page_num, "type": "table"},
                                chunk_id=str(uuid.uuid4())
                            ))

                # B. 提取原生文本
                text = page.extract_text()
                if text:
                    page_text_buffer += text + "\n"

                # C. [新增] 混合视觉解析 (图片 OCR)
                # 条件：已安装OCR库 + 开启增强模式 + 页面包含图片对象
                if HAS_OCR and self.use_advanced_mode and page.images:
                    try:
                        # 将整个页面渲染为图片 (resolution=200 兼顾速度与识别率)
                        # 这样做比提取单个图片更稳健，因为能保留图片排版位置的上下文
                        pil_image = page.to_image(resolution=200).original
                        ocr_text = self._run_ocr_on_image(pil_image)

                        if ocr_text:
                            # 简单去重：如果 OCR 识别出的内容很长，且不在原生文本里，则追加
                            # 或者直接追加，依靠 Semantic Splitter 去处理语义
                            if len(ocr_text) > 20: # 过滤太短的噪点
                                # 标记这部分内容来源于图片/扫描件
                                page_text_buffer += f"\n\n【第{page_num}页图片/扫描件内容】:\n{ocr_text}\n"
                                print(f"   👁️ [OCR] 第 {page_num} 页提取到图片文字 ({len(ocr_text)} 字符)")
                    except Exception as e:
                        print(f"   ⚠️ 第 {page_num} 页 OCR 处理失败: {e}")

                full_text_buffer += page_text_buffer

        if full_text_buffer:
            text_chunks = self._split_text(full_text_buffer)
            # 简单页码估算
            total_chars = len(full_text_buffer)
            chars_per_page = total_chars / total_pages if total_pages > 0 else 1000
            current_char_idx = 0
            for txt in text_chunks:
                chunk_len = len(txt)
                mid_point = current_char_idx + (chunk_len / 2)
                est_page = min(int(mid_point / chars_per_page) + 1, total_pages)
                current_char_idx += chunk_len
                all_chunks.append(IndexableChunk(
                    content=txt,
                    metadata={**self.project_info, "source_file": os.path.basename(file_path), "page": est_page, "type": "text"},
                    chunk_id=str(uuid.uuid4())
                ))
        return all_chunks

    def _split_text(self, text: str) -> List[str]:
        if self.use_advanced_mode:
            try:
                docs = self.splitter.create_documents([text])
                return [d.page_content for d in docs]
            except Exception as e:
                print(f"⚠️ 语义切分出错 ({e})，降级为递归切分")
                from langchain.text_splitter import RecursiveCharacterTextSplitter
                fallback = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
                return fallback.split_text(text)
        else:
            return self.splitter.split_text(text)
