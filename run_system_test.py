"""
run_system_test.py
功能：全链路系统集成测试 (End-to-End Integration Test)
覆盖模块：Ingestion -> ETL -> Retrieval -> Generation
"""

import os
import sys
import shutil
import time
from docx import Document
from typing import List

# 确保能导入各模块
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# 检查 API Key
if not os.getenv("DEEPSEEK_API_KEY"):
    print("❌ 错误: 未检测到 DEEPSEEK_API_KEY 环境变量。")
    print("请先执行: $env:DEEPSEEK_API_KEY='sk-xxxxxx' (PowerShell) 或 set DEEPSEEK_API_KEY=sk-xxx (CMD)")
    sys.exit(1)


class SystemIntegrityTester:
    def __init__(self):
        self.test_file_path = "temp_test_doc.docx"
        self.test_collection_name = "test_runner_collection"
        self.chunks = []
        print("\n🛡️  [系统自检] 开始全链路集成测试...\n" + "=" * 50)

    def _create_dummy_docx(self):
        """创建一个包含标题、文本和表格的临时 Word 文档"""
        doc = Document()
        doc.add_heading('第一章 土方开挖工程', level=1)
        doc.add_paragraph('1.1 施工准备', style='Heading 2')
        doc.add_paragraph('土方开挖前，应清理表层植被，清理厚度为 30cm。')
        doc.add_heading('1.2 机械配置', level=2)
        doc.add_paragraph('主要使用挖掘机和自卸汽车。具体配置如下表：')

        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "设备名称"
        table.cell(0, 1).text = "数量"
        table.cell(1, 0).text = "挖掘机 (PC200)"
        table.cell(1, 1).text = "4台"
        table.cell(2, 0).text = "自卸汽车"
        table.cell(2, 1).text = "10辆"

        doc.save(self.test_file_path)
        print(f"✅ [Setup] 生成临时测试文档: {self.test_file_path}")

    def step_1_ingestion(self):
        """测试解析层：Ingestion & Parsing"""
        print("\n📦 [Step 1] 测试解析层 (Ingestion)...")
        try:
            from ingestion.tender_parser import TenderDocParser

            self._create_dummy_docx()

            parser = TenderDocParser(project_info={"test_id": "001"})
            self.chunks = parser.parse_file(self.test_file_path)

            if len(self.chunks) > 0:
                print(f"   ✅ 解析成功! 生成了 {len(self.chunks)} 个切片。")
                # 验证 Parent-Child 逻辑
                parents = [c for c in self.chunks if c.is_parent]
                children = [c for c in self.chunks if not c.is_parent]
                print(f"   ℹ️  结构分析: Parent块={len(parents)}, Child块={len(children)}")

                # 验证表格是否转为 Markdown
                has_markdown_table = any("| 设备名称 |" in c.content for c in self.chunks)
                if has_markdown_table:
                    print("   ✅ 表格识别成功 (Markdown 格式检测通过)。")
                else:
                    print("   ⚠️  警告: 未检测到 Markdown 表格格式。")
            else:
                raise Exception("解析结果为空")

        except Exception as e:
            print(f"   ❌ 解析层测试失败: {e}")
            raise e

    def step_2_etl_vector_store(self):
        """测试存储层：Vector Store & ETL"""
        print("\n💾 [Step 2] 测试存储层 (Vector Store)...")
        try:
            from etl.vector_store import VectorStoreManager

            # 使用独立的测试集合，避免污染主库
            self.vs = VectorStoreManager(collection_name=self.test_collection_name)

            # 存入数据
            self.vs.add_chunks(self.chunks)

            # 简单验证存入数量
            count = self.vs.collection.count()
            if count > 0:
                print(f"   ✅ 数据入库成功! 当前测试库 chunk 数量: {count}")
            else:
                raise Exception("入库后数量为 0")

        except Exception as e:
            print(f"   ❌ 存储层测试失败: {e}")
            raise e

    def step_3_retrieval_modules(self):
        """测试检索层组件：Rewriter, Search, Compressor"""
        print("\n🔍 [Step 3] 测试检索层各组件 (Retrieval)...")
        try:
            # 1. 测试 Query Rewriter
            from retrieval.query_processor import QueryProcessor
            rewriter = QueryProcessor()
            original_q = "它需要几台挖掘机？"
            history = [{"role": "user", "content": "土方开挖的机械配置"}]
            rewritten_q = rewriter.rewrite(original_q, history)
            print(f"   ✅ [Rewriter] '{original_q}' -> '{rewritten_q}'")

            # 2. 测试 Search Engine (Search + Rerank)
            from retrieval.search_engine import TenderRetriever
            # 强制 Search Engine 使用我们的测试集合
            retriever = TenderRetriever()
            retriever.vector_db = self.vs  # 替换为刚才初始化的测试 DB 实例

            results = retriever.search("挖掘机数量", top_k=2)
            if results:
                print(f"   ✅ [Search] 检索成功，返回 {len(results)} 条结果。")
                print(f"      Top 1: {results[0]['content'][:30]}...")
            else:
                print("   ⚠️  [Search] 未检索到结果 (可能是数据太少被过滤或 Rerank 问题)。")

            # 3. 测试 Compressor
            from retrieval.compressor import ContextCompressor
            compressor = ContextCompressor()
            if results:
                compressed = compressor.compress("挖掘机数量", results)
                print(f"   ✅ [Compressor] 压缩完成 (长度: {len(compressed)} chars)。")

        except Exception as e:
            print(f"   ❌ 检索层测试失败: {e}")
            # 不阻断后续流程，只是标记失败
            pass

    def step_4_generation_rag(self):
        """测试生成层：RAG Service (End-to-End Chat)"""
        print("\n💬 [Step 4] 测试生成层 (RAG Chat)...")
        try:
            from generation.rag_service import DeepSeekRAGService

            service = DeepSeekRAGService()
            # 关键：Hack 一下，让 Service 使用我们的测试数据库
            service.retriever.vector_db = self.vs

            history = []
            question = "这个项目的土方开挖需要什么机械？"
            print(f"   🗣️  User: {question}")
            print("   🤖  AI Stream: ", end="")

            full_response = ""
            for event in service.chat_stream(question, history=history):
                if event['type'] == 'text':
                    print(event['data'], end="", flush=True)
                    full_response += event['data']
                elif event['type'] == 'status':
                    # 打印状态流
                    print(f"\n      [{event['data']}]", end="")
                elif event['type'] == 'error':
                    print(f"\n      ❌ Error: {event['data']}")

            print("\n")

            if len(full_response) > 10:
                print("   ✅ RAG 对话测试通过 (生成了有效回复)。")
            else:
                print("   ⚠️  RAG 回复过短，可能存在异常。")

        except Exception as e:
            print(f"   ❌ 生成层测试失败: {e}")

    def cleanup(self):
        """清理测试数据"""
        print("\n🧹 [Cleanup] 正在清理测试环境...")
        # 1. 删除临时文件
        if os.path.exists(self.test_file_path):
            os.remove(self.test_file_path)
            print(f"   - 已删除临时文件: {self.test_file_path}")

        # 2. 删除 Chroma 测试集合
        try:
            self.vs.client.delete_collection(self.test_collection_name)
            print(f"   - 已删除测试向量集合: {self.test_collection_name}")
        except:
            pass

        print("=" * 50 + "\n✅ 全链路测试结束。")


if __name__ == "__main__":
    tester = SystemIntegrityTester()
    try:
        tester.step_1_ingestion()
        tester.step_2_etl_vector_store()
        tester.step_3_retrieval_modules()
        tester.step_4_generation_rag()
    except KeyboardInterrupt:
        print("\n⛔ 测试被用户中断。")
    except Exception as e:
        print(f"\n⛔ 测试过程中发生严重错误: {e}")
    finally:
        tester.cleanup()
